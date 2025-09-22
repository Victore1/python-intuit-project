from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

#profile picture
document.add_picture('foto profil.jpg', width=Inches(1.0))

# name phone number and email details
name = input('What is your name? ')
speak('Hello ' + name + 'how are you today? ')
phone_number = input('What is your phone_number? ')
email = input('What is your email? ')
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('About me')
document.add_paragraph(
     input('Tell me about yourself? ')
)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph

company = input('Enter company name ') 
from_date = input ('From Date ')
to_date = input('To Date ')

p = document.add_paragraph()
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Decribe your experience ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph

        company = input('Enter company name ') 
        from_date = input ('From Date ')
        to_date = input('To Date ')

        p = document.add_paragraph()
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Decribe your experience ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter Skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Education History
document.add_heading('Education Experience')
p = document.add_paragraph

school = input('Enter school name ') 
from_date = input ('From Date ')
to_date = input('To Date ')

p = document.add_paragraph()
p.add_run(school + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

field_of_study = input(
    'Field of study ' + school)
p.add_run(field_of_study)

# more education history
while True:
    has_more_education_history = input(
        'Do you have more education history? Yes or No ')
    if has_more_education_history.lower() == 'yes':
        p = document.add_paragraph

        school = input('Enter school name ') 
        from_date = input ('From Date ')
        to_date = input('To Date ')

        p = document.add_paragraph()
        p.add_run(school + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        field_of_study = input(
            'Field of study ' + school + ' ')
        p.add_run(field_of_study)
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using KachTech course project"

document.save('cv.docx')